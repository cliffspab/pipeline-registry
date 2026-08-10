"""Build the full BKP compiled governance set from Markdown.

The Markdown file is the sole content authority.  This builder uses Pandoc only
as a Markdown parser, then creates native Word paragraphs, headings, lists,
tables, hyperlinks, fields and page furniture with python-docx.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import yaml

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Twips, Inches, Pt, RGBColor

from bkp_docx_design import (
    add_numbering_definition,
    add_page_field,
    apply_numbering,
    clear_document_body,
    set_cell_margins,
    set_cell_shading,
    set_font,
    set_paragraph_bottom_rule,
    set_paragraph_left_rule,
    set_repeat_table_layout,
    set_table_borders,
    set_update_fields,
)


# 080826: four parts became two. PROCESSES is a section inside CORE;
# STATUS and REFERENCES are branches of REGISTER.
COMPONENTS = ["CORE", "DIRECTORY"]

# A code block longer than this cannot be held on one page, so keep_together is
# not applied to it. Roughly a full page of Courier New 8.8 at these margins.
KEEP_TOGETHER_MAX_LINES = 40

# The strip is the reader's map and it has always had FOUR cells. It is not the
# same thing as COMPONENTS, which is the FILE structure: two parts, because the
# register must travel as one fenced YAML document. The content map did not
# merge — CORE still carries core and processes, REGISTER still carries status
# and references. Reducing the strip to two cells on the 080826 merge threw away
# half the map and left a two-cell indicator carrying almost no information.
#
# Cell width 9360/4 = 2340 twips. Position alone says which section this is:
# CORE hard left, PROCESSES a quarter across, STATUS half, REFS flush right.
STRIP = ["CORE", "PROCESSES", "STATUS", "REFERENCES"]

# Which strip cells each part owns. The part's title sits over the first of them.
PART_CELLS = {
    "CORE": ("CORE", "PROCESSES"),
    "DIRECTORY": ("STATUS", "REFERENCES"),
}

# The volume presents FOUR sections, one per strip cell, as it always has. The
# file carries two parts because the register must travel as one fenced YAML
# document; that is a transport constraint and it never described the reading
# order. PROCESSES opens on its H2 inside CORE; STATUS and REFERENCES are the
# register's two branches and open as the register is rendered.
SECTION_SUBTITLE = {
    "CORE": "what we do",
    "PROCESSES": "how we do it",
    "STATUS": "people in the news",
    "REFERENCES": "the knowledge base",
}

# REFERENCES prints as REFS: at Heading 1 size, indented to the fourth cell,
# the full word wraps. Presentation only - canon keeps the full name.
SECTION_DISPLAY = {"REFERENCES": "REFS"}

# The running head names where the text actually lives, which is the file the
# reader would fetch. PROCESSES is inside CORE; both register branches are /reg.
SECTION_SHORTFORM = {
    "CORE": "/core", "PROCESSES": "/core",
    "STATUS": "/reg", "REFERENCES": "/reg",
}

# YAML key -> the heading the original volume printed. Anything absent here is
# title-cased from its key, so a new branch still renders rather than vanishing.
REGISTER_HEADINGS = {
    "apex": "APEX FIGURES",
    "cabinet": "PORTFOLIO SWAPPERS",
    "reversals": "POLITICAL AND LEGAL REVERSALS",
    "mortalities": "MORTALITIES AND SUCCESSIONS",
    "corporate": "CORPORATE LEADERSHIP CHANGES",
    "global": "GLOBAL FIGURES",
    "countries": "COUNTRIES",
    "preferred_forms": "PREFERRED FORMS",
    "demonyms": "ADJECTIVAL / DEMONYM TRAPS",
    "structure": "CAPITALS AND COUNTRY STRUCTURE",
    "headline_abbreviations": "HEADLINE ABBREVIATION FORMS",
    "the_article": "THE ARTICLE",
    "foreign_places": "FOREIGN PLACES",
    "nomenclature": "NOMENCLATURE",
    "thai_places": "THAI PLACES",
    "transliteration_rules": "TRANSLITERATION MASTER RULES",
    "airports": "AIRPORTS",
    "third_party_spellings": "THIRD-PARTY SPELLINGS",
    "provinces": "PROVINCES",
    "organisations": "ORGANISATIONS",
    "vocabulary": "VOCABULARY & SPELLING",
    "numbers_symbols": "NUMBERS AND SYMBOLS",
    "uk_us_traps": "UK VS US TRAPS",
    "us_forms_acceptable": "US FORMS ACCEPTABLE",
    "rulings": "VOCABULARY RULINGS",
    "chinese": "CHINESE", "indonesian": "INDONESIAN",
    "myanmar_cambodian": "MYANMAR / CAMBODIAN",
    "given_name_as_surname": "GIVEN NAME AS SURNAME",
    "arabic_prefix": "ARABIC AL- PREFIX", "korean": "KOREAN",
    "vietnamese": "VIETNAMESE", "japanese": "JAPANESE",
    "icelandic": "ICELANDIC", "spanish": "SPANISH",
    "single_name_no_honorific": "SINGLE NAME, NO HONORIFIC",
    "confusable": "CONFUSABLE NAMES", "spelling_traps": "SPELLING TRAPS",
}

# Composed from the listed fields, in this order, for list-of-record branches.
RECORD_FIELDS = ["fact", "office", "ruling", "second_ref", "directive"]

# Running-head short forms - the shortlink slugs, so the page head reads the
# same way the operator addresses the section.
# Printed head for each component. REFERENCES prints as REFS: at Heading 1
# size, indented to the fourth register cell, "REFERENCES" wraps to two lines.
# The component keeps its full name - the split regex, part detection and the
# register all key off COMPONENTS - so this is presentation only and never
# reaches back into canon.
COMPONENT_DISPLAY = {
    "CORE": "CORE",
    "DIRECTORY": "DIRECTORY",
}

COMPONENT_SHORTFORM = {
    "CORE": "/core",
    "DIRECTORY": "/reg",
}
COMPONENT_DESCRIPTIONS = {
    "CORE": "Doctrine, authority, conventions, editing and output",
    "DIRECTORY": "Current tripwires, canonical forms and exceptions",
}
SEPARATOR_RE = re.compile(r"^={20,}$")


def ensure_style(doc, name, base="normal"):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
        return style


def configure_full_styles(doc):
    normal = doc.styles["normal"]
    set_font(normal, "Arial", 10.5)
    normal.paragraph_format.space_after = Pt(6.5)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_font(h1, "Arial Black", 46, True)   # section head, up from 36
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False

    h2 = doc.styles["Heading 2"]
    set_font(h2, "Arial Black", 20, True)
    h2.font.all_caps = True
    # Much lighter: at full black this level competed with the section head
    # and blunted it in thumbnail view. Tone separates them, size alone did not.
    h2.font.color.rgb = RGBColor(0xB0, 0xB0, 0xB0)
    # 270726: no rule at this level. The rule was the heaviest mark on the
    # line and made the grey heading read as a caption to it. Separation is
    # done with white space - roughly three blank lines above.
    h2.paragraph_format.space_before = Pt(46)

    # The reference document defines a bottom border on Heading 1 and
    # Heading 2. Clear it at style level or it reappears regardless of what
    # the paragraph does.
    for sid in ("Heading 1", "Heading 2"):
        el = doc.styles[sid]._element
        pPr = el.find(qn("w:pPr"))
        if pPr is not None:
            for bdr in pPr.findall(qn("w:pBdr")):
                pPr.remove(bdr)
    h2.paragraph_format.space_before = Pt(17)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_font(h3, "Arial Black", 13.2, True)
    h3.font.all_caps = True
    h3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    h3.paragraph_format.space_before = Pt(11)
    h3.paragraph_format.space_after = Pt(4.5)
    h3.paragraph_format.keep_with_next = True

    h4 = doc.styles["Heading 4"]
    set_font(h4, "Arial", 11.2, True)
    h4.paragraph_format.space_before = Pt(9)
    h4.paragraph_format.space_after = Pt(3.5)
    h4.paragraph_format.keep_with_next = True

    roles = {
        "BKP Cover Kicker": ("Arial Black", 14, True, False),
        "BKP Cover Display": ("Arial Black", 66, True, False),
        "BKP Cover Year": ("Arial Black", 60, True, False),
        "BKP Cover Tagline": ("Arial", 14, True, True),
        "BKP Part Stamp": ("Arial", 10, False, False),
        "BKP Cover Meta": ("Arial", 10, False, False),
        "BKP Register": ("Arial Black", 7.6, True, False),
        "BKP Part Label": ("Arial Black", 8.5, True, False),
        "BKP Metadata": ("Arial", 10, False, False),
        "BKP Part Subtitle": ("Arial", 10.5, False, False),
        "BKP Axiom Label": ("Arial Black", 7.3, True, False),
        "BKP Axiom": ("Arial", 12.2, True, False),
        "BKP Operator Note": ("Arial", 10, False, True),
        "BKP Code": ("Courier New", 8.8, False, False),
        "BKP Inline Code": ("Courier New", 9.5, False, False),
        "BKP Numbered": ("Arial", 10.2, False, False),
        "BKP Bullet": ("Arial", 10.2, False, False),
        "BKP List Continuation": ("Arial", 10.2, False, False),
        "BKP Footer": ("Arial", 7.8, True, False),
        "BKP Contents Label": ("Arial Black", 8, True, False),
        "BKP Contents Title": ("Arial Black", 15, True, False),
        "BKP Contents Description": ("Arial", 9.5, False, False),
        "BKP Source Marker": ("Courier New", 8.5, False, False),
        "BKP Status Subtitle": ("Arial", 10.5, True, False),
        "BKP Register Group": ("Arial Black", 10.5, True, False),
    }
    for name, (font, size, bold, italic) in roles.items():
        style = ensure_style(doc, name)
        set_font(style, font, size, bold, italic, "000000" if name != "BKP Footer" else "555555")

    doc.styles["BKP Cover Kicker"].paragraph_format.space_after = Pt(22)
    doc.styles["BKP Cover Display"].paragraph_format.space_after = Pt(0)
    doc.styles["BKP Cover Display"].paragraph_format.keep_with_next = True
    doc.styles["BKP Cover Year"].paragraph_format.space_before = Pt(8)
    doc.styles["BKP Cover Year"].paragraph_format.keep_with_next = True
    doc.styles["BKP Cover Year"].paragraph_format.space_after = Pt(18)
    doc.styles["BKP Cover Tagline"].paragraph_format.space_before = Pt(16)
    doc.styles["BKP Part Stamp"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.styles["BKP Part Stamp"].paragraph_format.space_before = Pt(2)
    doc.styles["BKP Part Stamp"].paragraph_format.space_after = Pt(10)
    doc.styles["BKP Cover Meta"].paragraph_format.space_after = Pt(3)
    doc.styles["BKP Register"].paragraph_format.space_after = Pt(0)
    doc.styles["BKP Part Label"].paragraph_format.space_before = Pt(10)
    doc.styles["BKP Part Label"].paragraph_format.space_after = Pt(5)
    doc.styles["BKP Metadata"].paragraph_format.space_after = Pt(4.5)
    # Tucked hard under the title: the Heading 1 above it already carries
    # 12pt after, so this pulls up against it rather than floating.
    # Body text in every respect - Arial 10.5, black, regular. It is a caption
    # to the title, not a display line, so it takes no treatment of its own.
    doc.styles["BKP Part Subtitle"].paragraph_format.space_before = Pt(0)
    doc.styles["BKP Part Subtitle"].paragraph_format.space_after = Pt(14)
    doc.styles["BKP Part Subtitle"].paragraph_format.line_spacing = 1.1
    doc.styles["BKP Part Subtitle"].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    doc.styles["BKP Axiom Label"].paragraph_format.space_before = Pt(7)
    doc.styles["BKP Axiom Label"].paragraph_format.space_after = Pt(1)
    doc.styles["BKP Axiom Label"].paragraph_format.left_indent = Inches(0.16)
    doc.styles["BKP Axiom Label"].paragraph_format.keep_with_next = True
    doc.styles["BKP Axiom"].paragraph_format.space_after = Pt(8)
    doc.styles["BKP Axiom"].paragraph_format.left_indent = Inches(0.16)
    doc.styles["BKP Operator Note"].paragraph_format.left_indent = Inches(0.18)
    doc.styles["BKP Operator Note"].paragraph_format.space_before = Pt(7)
    doc.styles["BKP Operator Note"].paragraph_format.space_after = Pt(7)
    doc.styles["BKP Numbered"].paragraph_format.space_after = Pt(4.5)
    doc.styles["BKP Numbered"].paragraph_format.line_spacing = 1.06
    doc.styles["BKP Bullet"].paragraph_format.space_after = Pt(4)
    doc.styles["BKP Bullet"].paragraph_format.line_spacing = 1.06
    doc.styles["BKP List Continuation"].paragraph_format.left_indent = Inches(0.34)
    doc.styles["BKP List Continuation"].paragraph_format.space_after = Pt(4)
    doc.styles["BKP Footer"].paragraph_format.space_after = Pt(0)
    doc.styles["BKP Contents Label"].paragraph_format.space_after = Pt(1)
    doc.styles["BKP Contents Title"].paragraph_format.space_after = Pt(1)
    doc.styles["BKP Contents Description"].paragraph_format.space_after = Pt(0)
    doc.styles["BKP Source Marker"].paragraph_format.space_after = Pt(5)
    doc.styles["BKP Status Subtitle"].paragraph_format.space_after = Pt(8)
    doc.styles["BKP Register Group"].font.all_caps = True
    doc.styles["BKP Register Group"].paragraph_format.space_before = Pt(8)
    doc.styles["BKP Register Group"].paragraph_format.space_after = Pt(4)
    doc.styles["BKP Register Group"].paragraph_format.keep_with_next = True
    h3.paragraph_format.space_after = Pt(6)
    doc.styles["BKP Bullet"].paragraph_format.space_before = Pt(1.5)
    doc.styles["BKP Numbered"].paragraph_format.space_before = Pt(1.5)


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.34)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_field(paragraph, instruction, cached_text=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = cached_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def configure_header_footer(section, label, blank_first=False):
    section.different_first_page_header_footer = blank_first
    short = COMPONENT_SHORTFORM.get(label, label)
    # The running head sits at the same horizontal position its cell occupies
    # in the four-bar register: CORE far left, PROCESSES a quarter across,
    # STATUS half, REFERENCES three-quarters. Placement alone says which
    # section you are in, so the head echoes the black box rather than
    # repeating it. Cell width is 2340 twips (the register grid).
    # Reverted 270726: the running head is pinned right, not stepped across
    # the page. Justifying it to the register cell was my misreading of the
    # head-position instruction, which applied to the section heads only.
    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        clear_paragraph(p)
        p.style = "BKP Footer"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Section slug washed out; the folio hard black so it reads as a page
        # number, not as part of the section label. Kevin needs to be able to
        # say "page 27" rather than "part 3, five paragraphs down".
        r = p.add_run(short + "  ")
        r.font.color.rgb = RGBColor(0xA8, 0xA8, 0xA8)
        r.font.size = Pt(10)
        add_page_field(p, bold=True, tab=False, size=Pt(10))
        # 090826: the rule under the running head is gone, op-ruled. It ran on
        # every page and under the version line on part openings, where it sat
        # between the stamp and the black block and separated two things that
        # were not in conflict. The washed-out slug and the hard black folio
        # already read as a head; a line under them was doing no work.
    # 270726: footer retired entirely - the folio moved into the running head.
    for footer in (section.footer, section.even_page_footer):
        footer.is_linked_to_previous = False
        clear_paragraph(footer.paragraphs[0])


def add_register(doc, active):
    # 080826: the grid derives from COMPONENTS rather than assuming four
    # cells. 9360 twips is the text width; it divides evenly by two and by
    # four, so a part-count change no longer needs this edited.
    cols = len(STRIP)
    table = doc.add_table(rows=1, cols=cols)
    widths = [9360 // cols] * cols
    set_repeat_table_layout(table, widths)
    for idx, name in enumerate(STRIP):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 55, 90, 55, 90)
        p = cell.paragraphs[0]
        p.style = doc.styles["BKP Register"]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if name == active:
            # 270726: the active cell renders as a solid black block with no
            # text - the page title already names the component. One cell
            # blacks at a time; the volume has four sections and each opens on
            # its own.
            set_cell_shading(cell, "000000")
            continue
        # Full names, caps, no slash and no ordinal. The slash form belongs to
        # pagination - it appears in the running head, where the folio needs a
        # compact label. REFERENCES stays fully titled here; REFS is used only
        # in the section head, where the long form will not fit its cell.
        p.add_run(name)
    set_table_borders(table, "000000", 10)
    row_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_pr.append(repeat)
    # 090826: was an empty paragraph at body size, which cost a full line
    # between the strip and the part title. Held as a hairline instead - the
    # table needs something after it, but not a line's worth.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(4)
    spacer.add_run().font.size = Pt(4)


def add_cover(doc, source_title, source_meta):
    p = doc.add_paragraph("A DETERMINISTIC, AI-FIRST REIMAGINING OF\nNEWSPAPER SUB-EDITING:", style="BKP Cover Kicker")
    p.paragraph_format.space_before = Pt(24)
    for line in ("THE", "BANGKOK", "POST", "BLUEPRINT"):
        doc.add_paragraph(line, style="BKP Cover Display")
    doc.add_paragraph("2026", style="BKP Cover Year")
    rule = doc.add_paragraph()
    set_paragraph_bottom_rule(rule, 34, 4)
    # 310726: the rule and tagline were free-floating, so on real Arial Black
    # metrics they broke to a second page and the tagline sat alone. Bound to
    # the display block; the 55pt lead padding was cut to Pt(24) to pay for it.
    rule.paragraph_format.keep_with_next = True
    doc.add_paragraph("This is why we have style...", style="BKP Cover Tagline")
    # 270726: source title and the slug/components line no longer print on the
    # cover. The slug still identifies the build - it is written to the Word
    # document properties instead (see build()), so the edition is recoverable
    # from File > Info without appearing on the page.


def add_contents(doc):
    title = doc.add_paragraph("INSTITUTIONAL REGISTER", style="Heading 1")
    title.paragraph_format.page_break_before = True
    set_paragraph_bottom_rule(title, 28, 5)
    p = doc.add_paragraph(
        "One maintained Markdown source. Four governed components. One mechanically generated Word document."
    )
    p.paragraph_format.space_after = Pt(18)
    table = doc.add_table(rows=4, cols=2)
    set_repeat_table_layout(table, [1580, 7780])
    for idx, name in enumerate(COMPONENTS):
        left, right = table.rows[idx].cells
        for cell in (left, right):
            set_cell_margins(cell, 135, 135, 135, 135)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, "000000")
        lp = left.paragraphs[0]
        lp.style = doc.styles["BKP Contents Label"]
        lr = lp.add_run(f"{idx + 1:02d}")
        lr.font.color.rgb = RGBColor(255, 255, 255)
        rp = right.paragraphs[0]
        rp.style = doc.styles["BKP Contents Title"]
        rp.add_run(name)
        desc = right.add_paragraph(COMPONENT_DESCRIPTIONS[name], style="BKP Contents Description")
        if idx % 2:
            set_cell_shading(right, "F2F2F2")
    set_table_borders(table, "000000", 10)
    first_row_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    first_row_pr.append(repeat)
    p = doc.add_paragraph("BUILD CONTRACT", style="Heading 2")
    p.paragraph_format.space_before = Pt(20)
    note = doc.add_paragraph(style="BKP Operator Note")
    set_paragraph_left_rule(note, 18, 8)
    r = note.add_run("SOURCE OF TRUTH  ")
    r.bold = True
    r.italic = False
    note.add_run(
        "The compiled Markdown governs wording, order and semantic structure. "
        "Word supplies the repeatable visual and navigational system."
    )


def restart_page_numbering(section, start=1):
    """Page numbers restart at 1 in each part.

    Section numbering, not continuous: with a single edit surface rebuilt on
    every push, continuous folios make citations perishable - adding a
    paragraph to CORE shifts every page number after it. Restarting per part
    confines that churn, and the running head names the section, so "core 7"
    is unambiguous and survives being quoted without context.
    """
    sectPr = section._sectPr
    existing = sectPr.find(qn("w:pgNumType"))
    if existing is not None:
        sectPr.remove(existing)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(start))
    # Schema order: pgNumType sits after pgBorders/lnNumType, before cols.
    cols = sectPr.find(qn("w:cols"))
    if cols is not None:
        cols.addprevious(pg)
    else:
        sectPr.append(pg)


def add_part_opening(doc, component, part_number, subtitle=None):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(section)
    # blank_first: the lead page of each part already carries the register
    # strip and the component title, so the running head would only repeat it.
    configure_header_footer(section, component, blank_first=True)
    add_register(doc, component)
    # 270726: the "PART: X" label is gone - the title states it and the
    # register strip shows position. Two "part" statements on one page was
    # the redundancy.
    # The head sits over its cell in the register strip: CORE hard left,
    # PROCESSES a quarter across, STATUS half, REFS flush right. Position
    # alone says which section this is, echoing the blank black cell above
    # and the running head on every page that follows. Cell width 2340 twips.
    title = doc.add_paragraph(
        SECTION_DISPLAY.get(component, COMPONENT_DISPLAY.get(component, component)),
        style="Heading 1")
    # The title sits over its own cell. 090826 — this keyed off COMPONENTS,
    # which had become two, so REGISTER counted as the LAST cell and anchored
    # flush right, over REFERENCES rather than over anything it owned. It now
    # keys off the strip, which is where the cells actually are.
    i = STRIP.index(component) if component in STRIP else None
    if i is not None:
        if i == len(STRIP) - 1:
            # Last cell shares the right margin, so anchor to that edge.
            title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif i > 0:
            title.paragraph_format.left_indent = Twips(i * 2340)
    # STATUS and REFERENCES come out of YAML, which carries no sub-line of its
    # own; CORE and PROCESSES take theirs from the source text as before.
    if subtitle:
        # 100826, op-requested: the sub-line sits UNDER ITS TITLE, not adrift at
        # the left margin. It was a plain paragraph, so REFS anchored flush right
        # and "the knowledge base" started three columns away from it. It now
        # takes the title's own horizontal position - same indent, same
        # alignment - so the two read as one block.
        # The gap is the Heading 1's own space_after, 12pt. Zeroed on THIS
        # title only, so the sub-line sits directly beneath the letter above.
        title.paragraph_format.space_after = Pt(0)
        sub = doc.add_paragraph(subtitle, style="BKP Part Subtitle")
        sub.alignment = title.alignment
        sub.paragraph_format.left_indent = title.paragraph_format.left_indent
    # 270726: no rule under the head. The shortlink and version line sit
    # directly beneath it, and the rule moves below that pair - see the
    # slug post-pass in build().
    return section


def add_labeled_metadata(doc, text, labels):
    positions = []
    for label in labels:
        pos = text.find(label)
        if pos >= 0:
            positions.append((pos, label))
    positions.sort()
    for idx, (pos, label) in enumerate(positions):
        start = pos + len(label)
        end = positions[idx + 1][0] if idx + 1 < len(positions) else len(text)
        value = text[start:end].strip()
        p = doc.add_paragraph(style="BKP Metadata")
        r = p.add_run(label)
        r.bold = True
        if value:
            p.add_run(" " + value)


def split_on_softbreak(inlines):
    segments = [[]]
    for item in inlines:
        if item["t"] == "SoftBreak":
            segments.append([])
        else:
            segments[-1].append(item)
    return [segment for segment in segments if inlines_text(segment).strip()]


def add_segmented_paragraph(doc, inlines):
    for segment in split_on_softbreak(inlines):
        text = inlines_text(segment).strip()
        if text.startswith("provinces:"):
            add_province_index(doc, text)
        elif re.fullmatch(r"--- .+ ---", text):
            p = doc.add_paragraph(style="BKP Register Group")
            add_inlines(p, segment)
            set_paragraph_bottom_rule(p, 6, 2)
        elif re.fullmatch(r"[a-z_]+:", text) or re.fullmatch(r"[A-Z][A-Z0-9_]+", text):
            p = doc.add_paragraph(style="BKP Source Marker")
            add_inlines(p, segment)
        else:
            match = re.match(r"^([A-Z][A-Z0-9 &'()./-]+:)\s*(.*)$", text)
            p = doc.add_paragraph(style="normal")
            if match:
                r = p.add_run(match.group(1))
                r.bold = True
                if match.group(2):
                    p.add_run(" " + match.group(2))
            else:
                add_inlines(p, segment)


def add_province_index(doc, text):
    marker, payload = text.split(":", 1)
    p = doc.add_paragraph(style="BKP Source Marker")
    p.add_run(marker + ":")
    payload = payload.strip()
    starts = [
        match.start(1)
        for match in re.finditer(
            r"(?:^|(?<=\]) )([A-Z][A-Za-z'-]*(?: [A-Z][A-Za-z'-]*)*): \[",
            payload,
        )
    ]
    records = [
        payload[start : starts[index + 1] if index + 1 < len(starts) else len(payload)].strip()
        for index, start in enumerate(starts)
    ]
    for record in records:
        match = re.match(r"^([^:]+:)\s*(.*)$", record)
        p = doc.add_paragraph(style="normal")
        p.paragraph_format.keep_together = True
        if match:
            label = p.add_run(match.group(1))
            label.bold = True
            p.add_run(" " + match.group(2))
        else:
            p.add_run(record)


def inlines_text(inlines):
    out = []
    for item in inlines:
        kind = item["t"]
        content = item.get("c")
        if kind == "Str":
            out.append(content)
        elif kind in ("Space", "SoftBreak", "LineBreak"):
            out.append(" " if kind != "LineBreak" else "\n")
        elif kind in ("Strong", "Emph", "Underline", "Strikeout", "SmallCaps", "Superscript", "Subscript"):
            out.append(inlines_text(content))
        elif kind == "Code":
            out.append(content[1])
        elif kind == "Link":
            out.append(inlines_text(content[1]))
        elif kind == "Image":
            out.append(inlines_text(content[1]))
        elif kind == "Quoted":
            quote = '"' if content[0]["t"] == "DoubleQuote" else "'"
            out.append(quote + inlines_text(content[1]) + quote)
        elif kind == "RawInline":
            out.append(content[1])
        elif kind == "Note":
            out.append("")
    return "".join(out)


def add_hyperlink(paragraph, text, url, bold=False, italic=False):
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    # 270726: shortlinks print bare - no underline, grey, bolder - so the
    # line matches the version/slug line beneath it and the two read as a pair.
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "808080")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rpr.extend([color, underline])
    rpr.append(OxmlElement("w:b"))
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inlines(paragraph, inlines, bold=False, italic=False):
    for item in inlines:
        kind = item["t"]
        content = item.get("c")
        if kind == "Str":
            run = paragraph.add_run(content)
            run.bold = bold
            run.italic = italic
        elif kind == "Space":
            paragraph.add_run(" ")
        elif kind == "SoftBreak":
            paragraph.add_run(" ")
        elif kind == "LineBreak":
            paragraph.add_run().add_break()
        elif kind == "Strong":
            add_inlines(paragraph, content, True, italic)
        elif kind == "Emph":
            add_inlines(paragraph, content, bold, True)
        elif kind == "Underline":
            start = len(paragraph.runs)
            add_inlines(paragraph, content, bold, italic)
            for run in paragraph.runs[start:]:
                run.underline = True
        elif kind == "Strikeout":
            start = len(paragraph.runs)
            add_inlines(paragraph, content, bold, italic)
            for run in paragraph.runs[start:]:
                run.font.strike = True
        elif kind == "Code":
            run = paragraph.add_run(content[1])
            run.bold = bold
            run.italic = italic
            run.font.name = "Courier New"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(9.5)
        elif kind == "Link":
            add_hyperlink(paragraph, inlines_text(content[1]), content[2][0], bold, italic)
        elif kind == "Quoted":
            quote = '"' if content[0]["t"] == "DoubleQuote" else "'"
            paragraph.add_run(quote)
            add_inlines(paragraph, content[1], bold, italic)
            paragraph.add_run(quote)
        elif kind == "Superscript":
            start = len(paragraph.runs)
            add_inlines(paragraph, content, bold, italic)
            for run in paragraph.runs[start:]:
                run.font.superscript = True
        elif kind == "Subscript":
            start = len(paragraph.runs)
            add_inlines(paragraph, content, bold, italic)
            for run in paragraph.runs[start:]:
                run.font.subscript = True
        elif kind == "RawInline":
            paragraph.add_run(content[1])


def block_text(block):
    kind = block["t"]
    content = block.get("c")
    if kind in ("Para", "Plain"):
        return inlines_text(content)
    if kind == "Header":
        return inlines_text(content[2])
    if kind == "CodeBlock":
        return content[1]
    if kind in ("BulletList", "OrderedList"):
        items = content if kind == "BulletList" else content[1]
        return " ".join(block_text(b) for item in items for b in item)
    return ""


def add_axiom(doc, inlines):
    label = doc.add_paragraph("AXIOM", style="BKP Axiom Label")
    set_paragraph_left_rule(label, 28, 9)
    p = doc.add_paragraph(style="BKP Axiom")
    set_paragraph_left_rule(p, 28, 9)
    add_inlines(p, inlines)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    set_paragraph_bottom_rule(p, 8, 2)


def set_paragraph_box(paragraph, fill="F4F4F4"):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "10")
        node.set(qn("w:space"), "7")
        node.set(qn("w:color"), "000000")
        borders.append(node)


def add_code_block(doc, text):
    p = doc.add_paragraph(style="BKP Code")
    p.style = doc.styles["BKP Code"]
    p.paragraph_format.left_indent = Inches(0.06)
    p.paragraph_format.right_indent = Inches(0.06)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    lines = text.splitlines()
    # 090826: keep_together only where it can be honoured. The register is a
    # single fenced block of 600+ lines - one paragraph, one break per line -
    # and asking Word to keep that on one page cannot succeed. It responds by
    # pushing the whole paragraph to a fresh page and splitting it there
    # anyway, which left the REGISTER part opening as a near-empty page with
    # the title stranded at the top. A block that cannot fit a page must be
    # allowed to break where it falls.
    p.paragraph_format.keep_together = len(lines) <= KEEP_TOGETHER_MAX_LINES
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        p.add_run(line)
    set_paragraph_box(p)


def register_heading(key):
    return REGISTER_HEADINGS.get(key, key.replace("_", " ").upper())


def add_entry(doc, key, value):
    """One register entry: bold key, en-dash, ruling. The 310726 sweep put the
    register on this shape - key + ruling - and converted the old tables to it.
    It is reproduced here rather than reverted to tables."""
    p = doc.add_paragraph()
    run = p.add_run(str(key))
    run.bold = True
    p.add_run(" — " + str(value).strip())
    return p


def compose_record(record):
    """A list-of-records branch (apex, cabinet members, reversals, global)
    carries name plus some of fact/office/ruling/second_ref/directive. Join the
    ones present, in a fixed order, so the entry reads as one sentence."""
    parts = [str(record[f]).strip().rstrip(".")
             for f in RECORD_FIELDS if record.get(f)]
    return ". ".join(parts) + "." if parts else ""


def render_register_node(doc, node, level):
    """Walk a register branch and emit headings and entries.

    Mapping, which is the original volume's shape:
      dict of scalars      -> one entry per pair
      dict of containers   -> heading at this level, then recurse
      list of records      -> one entry per record, keyed on name
      list of scalars      -> bullets
    'note' and 'convention' keys are prose and print as body, not as entries.
    """
    heading_style = "Heading %d" % min(level, 4)

    if isinstance(node, dict):
        for key, value in node.items():
            if key in ("note", "convention") and not isinstance(value, (dict, list)):
                doc.add_paragraph(str(value).strip())
                continue
            if isinstance(value, dict):
                doc.add_paragraph(register_heading(key), style=heading_style)
                render_register_node(doc, value, level + 1)
            elif isinstance(value, list):
                if value and isinstance(value[0], dict):
                    doc.add_paragraph(register_heading(key), style=heading_style)
                    for record in value:
                        add_entry(doc, record.get("name", ""), compose_record(record))
                else:
                    # A short scalar list reads better inline than as bullets;
                    # a long one is a real list. Province rosters are the long
                    # case and there are 77 of them.
                    add_entry(doc, register_heading(key).title(),
                              "; ".join(str(v) for v in value))
            else:
                add_entry(doc, key, value)
    elif isinstance(node, list):
        for item in node:
            if isinstance(item, dict):
                add_entry(doc, item.get("name", ""), compose_record(item))
            else:
                doc.add_paragraph(str(item), style="BKP Bullet")


def render_register(doc, data, open_section):
    """The register, as document structure rather than a code panel.

    Until 090826 the whole register printed as one fenced code block: 618 lines
    of YAML in a fixed-width panel, so its forty-odd headings did not exist in
    the volume and the contents page stopped at CORE. The fence is required in
    the SOURCE - it is what keeps the register parseable and the split
    invertible - but it is a transport wrapper, and printing the wrapper was
    never the intent. Parsed here and rendered as the original did.
    """
    before = len(doc.paragraphs)
    for branch, title in (("status", "STATUS"), ("references", "REFERENCES")):
        if branch not in data:
            continue
        open_section(title)
        render_register_node(doc, data[branch], 2)

    # The coverage audit matched the register as ONE unit - the fenced block,
    # verbatim. Rendered as structure it no longer appears verbatim, so that
    # unit is declared absent in compile.yml. Declaring it and stopping would
    # leave the single largest piece of content in the corpus unguarded, so the
    # check moves here and gets finer: every leaf value in the register must
    # reach a paragraph. Counted, not sampled.
    leaves = count_register_leaves(data)
    written = len(doc.paragraphs) - before
    if written < leaves:
        raise SystemExit(
            "FATAL: register render dropped content. %d leaf values in the "
            "YAML, %d paragraphs written. Every entry must reach the page."
            % (leaves, written))
    print("register render: PASS - %d leaf values, %d paragraphs"
          % (leaves, written))


def count_register_leaves(node):
    """Every scalar that must reach the page. Prose keys count too - they are
    printed as body rather than as entries, but they are still content."""
    if isinstance(node, dict):
        return sum(count_register_leaves(v) for v in node.values())
    if isinstance(node, list):
        if node and isinstance(node[0], dict):
            return len(node)          # one composed entry per record
        return 1                      # a scalar list prints as one entry
    return 1


def table_rows(table_block):
    content = table_block["c"]
    head_rows = content[3][1]
    body_groups = content[4]
    rows = []
    for row in head_rows:
        rows.append(("head", row[1]))
    for body in body_groups:
        for row in body[3]:
            rows.append(("body", row[1]))
    return rows


def cell_inlines(cell):
    blocks = cell[4]
    merged = []
    for idx, block in enumerate(blocks):
        if idx:
            merged.append({"t": "LineBreak"})
        if block["t"] in ("Plain", "Para"):
            merged.extend(block["c"])
        else:
            merged.append({"t": "Str", "c": block_text(block)})
    return merged


def add_table(doc, table_block):
    rows_data = table_rows(table_block)
    if not rows_data:
        return
    col_count = len(rows_data[0][1])
    table = doc.add_table(rows=len(rows_data), cols=col_count)
    if col_count == 3:
        widths = [2850, 2920, 3590]
    elif col_count == 4:
        widths = [1900, 2480, 2680, 2300]
    else:
        widths = [9360 // col_count] * col_count
        widths[-1] += 9360 - sum(widths)
    set_repeat_table_layout(table, widths)
    set_table_borders(table, "000000", 8)
    for r_idx, (role, cells) in enumerate(rows_data):
        for c_idx, cell_data in enumerate(cells):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 105, 115, 105, 115)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.03
            if role == "head":
                set_cell_shading(cell, "000000")
                p.style = doc.styles["BKP Register"]
                start = len(p.runs)
                add_inlines(p, cell_inlines(cell_data), bold=True)
                for run in p.runs[start:]:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, "F3F3F3")
                p.style = doc.styles["normal"]
                for run in p.runs:
                    run.font.size = Pt(8.4)
                add_inlines(p, cell_inlines(cell_data))
                for run in p.runs:
                    run.font.size = Pt(8.4)
    first_row_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    first_row_pr.append(repeat)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(5)


def render_list(doc, block, decimal_num_id, bullet_num_id, level=0):
    kind = block["t"]
    items = block["c"] if kind == "BulletList" else block["c"][1]
    num_id = bullet_num_id if kind == "BulletList" else add_numbering_definition(doc, "decimal", "%1.")
    style = "BKP Bullet" if kind == "BulletList" else "BKP Numbered"
    for item in items:
        first_text = True
        for child in item:
            if child["t"] in ("Plain", "Para"):
                for segment in split_on_softbreak(child["c"]):
                    text = inlines_text(segment).strip()
                    if text.startswith("provinces:"):
                        add_province_index(doc, text)
                        first_text = False
                        continue
                    if re.fullmatch(r"--- .+ ---", text):
                        p = doc.add_paragraph(style="BKP Register Group")
                        add_inlines(p, segment)
                        set_paragraph_bottom_rule(p, 6, 2)
                        continue
                    p = doc.add_paragraph(style=style if first_text else "BKP List Continuation")
                    if first_text:
                        apply_numbering(p, num_id)
                    add_inlines(p, segment)
                    first_text = False
            elif child["t"] in ("BulletList", "OrderedList"):
                render_list(doc, child, decimal_num_id, bullet_num_id, level + 1)
            elif child["t"] == "CodeBlock":
                add_code_block(doc, child["c"][1])


# Component title as it appears in the source H1 -> canonical component name.
# REFS is the display head; REFERENCES is the component and register name.
# 110826: the SEAM words are the file names (GUIDE, DIRECTORY); the values
# are the internal component names the rest of this module keys on. The
# volume still has four parts - CORE, PROCESSES, STATUS, REFERENCES - and
# none of them move. This dict is also the set of source H1s to swallow,
# since add_part_opening prints the title itself.
COMPONENT_TITLES = {"GUIDE": "CORE", "DIRECTORY": "DIRECTORY"}

PART_SEAM_RE = re.compile(r"<!--\s*PART:\s*(\S+)\s+(\w+)\s*-->")

# The version/shortlink pair under each part title. Right-aligned in the
# volume per operator ruling 010826; plain text keeps it left, where column
# padding would break on any reflow.
PART_STAMP_RE = re.compile(r"^go\.fuzzylogic\.page/\w+\s*$")


def set_first_page_stamp(section, text):
    """Full edition stamp as the opening page's running head.

    010826: the stamp was a body paragraph at the top of each part, which
    LOOKED like a header in Word while the real first-page header sat empty.
    Putting it where it belongs means the header band carries the same kind of
    thing on every page of a part - full stamp on the opening page, compact
    "/core 7" thereafter - so the eye tracks one position instead of two.
    """
    header = section.first_page_header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.style = "BKP Footer"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0xA8, 0xA8, 0xA8)
    r.font.size = Pt(10)
    set_paragraph_bottom_rule(p, 6, 2)


def is_axiom_candidate(blocks, index, last_heading_level):
    if last_heading_level != 3 or index + 1 >= len(blocks):
        return False
    block = blocks[index]
    if block["t"] != "Para":
        return False
    text = block_text(block).strip()
    if not (4 <= len(text) <= 95 and text.endswith((".", ":", "—"))):
        return False
    # Register entries open with a bold key ("**Niger** — Nigerien, not Nigerian.").
    # They are lookup rows, not doctrine: AXIOM is reserved for prose that earns it.
    content = block.get("c") or []
    if content and content[0].get("t") == "Strong":
        return False
    return blocks[index + 1]["t"] in ("Para", "BulletList", "OrderedList")


def write_manifest(path, source, output, ast, skipped_separators):
    payload = {
        "source": str(source.resolve()),
        "output": str(output.resolve()),
        "pandoc_api_version": ast.get("pandoc-api-version"),
        "source_blocks": len(ast["blocks"]),
        "source_block_types": {
            kind: sum(1 for block in ast["blocks"] if block["t"] == kind)
            for kind in sorted({block["t"] for block in ast["blocks"]})
        },
        "suppressed_structural_separator_paragraphs": skipped_separators,
        "design_components": COMPONENTS,
    }
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def parse_markdown(pandoc, source):
    with tempfile.TemporaryDirectory(prefix="bkp-pandoc-") as tmp:
        ast_path = Path(tmp) / "ast.json"
        subprocess.run(
            [str(pandoc), "-f", "gfm", "-t", "json", str(source), "-o", str(ast_path)],
            check=True,
        )
        return json.loads(ast_path.read_text(encoding="utf-8"))


def add_contents(doc):
    """A contents list built as ordinary paragraphs, not a Word TOC field.

    CURRENTLY OFF - the call in build() is commented out. Op-ruled 100826: an
    index without page numbers is a list, not an index. Adding numbers needs a
    two-pass build, which is straightforward, but it is blocked on a design
    ruling first: page numbering RESTARTS in every part (restart_page_numbering
    below), so the folio reads "/core 2" and STATUS begins at 1 again. A
    page-numbered index cannot address that volume without either qualifying
    every entry by part or running the numbering continuously.

    RULED 100826: qualify by part. "REFS 1" is the wanted form, so the restart
    stays. To finish it:
      1. Uncomment the call. Keep the entry paragraphs at the SAME line count as
         the numbered version will have, so pass one paginates identically and
         the numbers measured stay true - the contents fits one page at 22
         entries.
      2. Build, convert to PDF, and read the folio off each page: every page
         already prints SECTION_SHORTFORM + number in its header, e.g.
         "/core 2". Match each heading to the first page it appears on AFTER
         the contents page.
      3. The part for each entry is the Heading 1 above it in the entries list -
         CORE, PROCESSES, STATUS, REFS - not the shortform, which collapses
         PROCESSES into /core and both branches into /reg.
      4. Rebuild with "PART n" appended to each entry, convert again.

    A TOC field is populated by Word on open and cached in the file. CI has no
    Word: it builds the volume and converts it to PDF in the same job, so the
    field is still empty at conversion and the published PDF gets no index.
    LibreOffice will not populate it either - tested 100826, both with the bare
    field and with w:updateFields set. Opening the docx by hand between the two
    steps would work and is exactly the kind of step nobody owns.

    So the list is content. Real paragraphs, internal links to bookmarks on the
    headings, no field anywhere. It converts like any other text, which means
    the docx and the PDF carry the same index. No page numbers - they would need
    a layout pass the builder does not have, and the volume is read on screen.
    """
    entries = []
    for para in doc.paragraphs:
        if para.style.name not in ("Heading 1", "Heading 2"):
            continue
        text = para.text.strip()
        if not text:
            continue
        entries.append((para.style.name, text, para))
    if not entries:
        return

    anchor = next((p for p in doc.paragraphs if p.style.name == "Heading 1"), None)
    if anchor is None:
        return

    for index, (_, _, para) in enumerate(entries):
        name = "_bkp_toc_%d" % index
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(9000 + index))
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(9000 + index))
        para._p.insert(0, start)
        para._p.append(end)

    block = []
    head = doc.add_paragraph("CONTENTS", style="Heading 2")
    block.append(head._p)
    for index, (style, text, _) in enumerate(entries):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.0 if style == "Heading 1" else 0.28)
        pf.space_after = Pt(2)
        link = OxmlElement("w:hyperlink")
        link.set(qn("w:anchor"), "_bkp_toc_%d" % index)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        if style == "Heading 1":
            rpr.append(OxmlElement("w:b"))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20" if style == "Heading 1" else "18")
        rpr.append(sz)
        run.append(rpr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        link.append(run)
        p._p.append(link)
        block.append(p._p)

    # The contents ends the front matter: CORE opens on a fresh page after it.
    brk = doc.add_paragraph()
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    brk._p.append(run)
    block.append(brk._p)

    for element in block:
        anchor._p.addprevious(element)


def build(source, reference, output, pandoc, manifest, component=None):
    ast = parse_markdown(pandoc, source)
    shutil.copy2(reference, output)
    doc = Document(output)
    clear_document_body(doc)
    configure_full_styles(doc)
    doc.settings.odd_and_even_pages_header_footer = True
    configure_page(doc.sections[0])
    configure_header_footer(doc.sections[0], "2026", blank_first=True)
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.")
    bullet_num_id = add_numbering_definition(doc, "bullet", "●")

    blocks = ast["blocks"]

    # --component scope: select content only. The cover, the four-part
    # register and every style/list/table/header path stay exactly as in the
    # full build; only the block range changes. Requires the PART header to
    # survive so add_part_opening() still fires and the register activates
    # the right cell.
    if component:
        component = component.strip().upper()
        if component not in COMPONENTS:
            raise SystemExit("unknown component %r; expected one of %s"
                             % (component, ", ".join(COMPONENTS)))
        head, start, end = blocks[:2], None, len(blocks)
        for i, b in enumerate(blocks):
            if b["t"] != "Header":
                continue
            lvl, _, inl = b["c"]
            if lvl != 1:
                continue
            txt = inlines_text(inl).strip().upper()
            if txt not in COMPONENT_TITLES:
                continue
            name = COMPONENT_TITLES[txt]
            if name == component:
                start = i
            elif start is not None:
                end = i
                break
        if start is None:
            raise SystemExit("component %r not found in source" % component)
        blocks = head + blocks[start:end]
        ast["blocks"] = blocks
    source_title = blocks[0]["c"][2] if blocks and blocks[0]["t"] == "Header" else []
    source_meta = blocks[1]["c"] if len(blocks) > 1 and blocks[1]["t"] == "Para" else []
    add_cover(doc, source_title, source_meta)
    # 270726: the INSTITUTIONAL REGISTER page (and the BUILD CONTRACT note that
    # sat on it) is removed. add_contents() is retained but uncalled.

    # Slug -> Word document properties, so the edition survives the cover edit.
    try:
        slug = inlines_text(source_meta).strip().split()[0] if source_meta else ""
        doc.core_properties.title = "The Bangkok Post Blueprint 2026"
        if slug:
            doc.core_properties.subject = slug
            doc.core_properties.comments = "Edition %s - generated from BLUEPRINT.txt" % slug
    except Exception:
        pass

    # 010826: was a hard 2 - title block, slug block, then body. Splitting the
    # slug and the Components line onto separate source lines made them two
    # blocks, so Components fell through and printed under the cover tagline.
    # The front matter is cover material by definition: everything before the
    # first part seam belongs to add_cover, however many paragraphs it runs to.
    start_index = 2
    for _i, _b in enumerate(blocks):
        if _b["t"] == "RawBlock" and PART_SEAM_RE.search((_b["c"][1] or "")):
            start_index = _i
            break
    part_number = 0
    current_edition = ""
    current_component = None
    skip_duplicate_component_heading = False
    last_heading_level = None
    skipped_separators = 0
    skip_indices = set()

    for index in range(start_index, len(blocks)):
        if index in skip_indices:
            continue
        block = blocks[index]
        kind = block["t"]
        if kind == "Para" and SEPARATOR_RE.fullmatch(block_text(block).strip()):
            skipped_separators += 1
            continue
        if kind == "Para" and PART_STAMP_RE.match(block_text(block).strip()):
            # Lifted into the opening page's running head, not printed in the
            # body. The source keeps the line: it is the component's identity
            # when the .txt is read on its own, and the edition guard's anchor.
            set_first_page_stamp(doc.sections[-1],
                                 "%s | %s" % (current_edition, block_text(block).strip()))
            last_heading_level = None
            continue
        # 010826: the part opens on the SEAM, not on the title. The seam is
        # the boundary; keying off the title stranded anything between the two
        # - the version/shortlink line now sits there - on the previous part's
        # last page. pandoc emits the HTML comment as RawBlock.
        if kind == "RawBlock":
            fmt, raw = block["c"]
            m = PART_SEAM_RE.search(raw or "")
            if m:
                current_edition = m.group(1)
                current_component = COMPONENT_TITLES[m.group(2).upper()]
                # 090826: the REGISTER seam no longer opens a part. The volume
                # has no REGISTER section - it has STATUS and REFERENCES, the
                # two branches the register carries - and those open as the
                # register is rendered, below. Opening here would put an empty
                # REGISTER page in front of them.
                if current_component == "DIRECTORY":
                    last_heading_level = None
                    continue
                part_number += 1
                add_part_opening(doc, current_component, part_number)
                last_heading_level = None
                continue
            continue
        if kind == "Header":
            level, _, inlines = block["c"]
            text = inlines_text(inlines).strip()
            # The component title is an ordinary H1 now: the seam above it has
            # already opened the part, and add_part_opening prints the name, so
            # printing it again would duplicate it.
            if level == 1 and text.upper() in COMPONENT_TITLES:
                last_heading_level = 1
                continue
            # 090826: PROCESSES is a SECTION of the file (an H2 inside CORE)
            # and a PART of the volume. The 080826 merge collapsed both at
            # once; only the file one had to move. It opens on its own heading.
            if level == 2 and text.upper() == "PROCESSES":
                part_number += 1
                add_part_opening(doc, "PROCESSES", part_number)
                last_heading_level = None
                continue
            if level == 3 and text.startswith("Verified Editorial Status Changes:"):
                doc.add_paragraph(text, style="BKP Status Subtitle")
                last_heading_level = None
                continue
            if level == 2 and text.startswith("Source:") and "As of:" in text:
                add_labeled_metadata(doc, text, ["Source:", "As of:", "House Style Note:"])
                last_heading_level = None
                continue
            style = f"Heading {min(level, 4)}"
            p = doc.add_paragraph(style=style)
            add_inlines(p, inlines)
            # 270726: no rule at either level. The section head carries its
            # rule under the shortlink/version pair instead, and the level
            # below is separated by white space, not by a line.
            last_heading_level = level
            continue
        if kind == "Para":
            text = block_text(block).strip()
            closing_tag = None
            if text == "DDMMYY — Slug":
                closing_tag = "</page_ready>"
            elif text == "<state_log>":
                closing_tag = "</state_log>"
            if closing_tag:
                gathered = []
                for inner_index in range(index, len(blocks)):
                    inner = blocks[inner_index]
                    if inner["t"] != "Para":
                        break
                    gathered.append(block_text(inner).strip())
                    skip_indices.add(inner_index)
                    if block_text(inner).strip() == closing_tag:
                        break
                add_code_block(doc, "\n\n".join(gathered))
                last_heading_level = None
                continue
            if is_axiom_candidate(blocks, index, last_heading_level):
                add_axiom(doc, block["c"])
            else:
                if all(label in text for label in ("Status:", "Architecture:", "Companions:")):
                    add_labeled_metadata(doc, text, ["Status:", "Architecture:", "Companions:"])
                elif all(label in text for label in ("Status:", "Purpose:")):
                    add_labeled_metadata(doc, text, ["Status:", "Purpose:"])
                elif text.startswith("provinces:"):
                    add_province_index(doc, text)
                elif any(item["t"] == "SoftBreak" for item in block["c"]):
                    add_segmented_paragraph(doc, block["c"])
                else:
                    style = "BKP Source Marker" if len(block["c"]) == 1 and block["c"][0]["t"] == "Code" else "normal"
                    p = doc.add_paragraph(style=style)
                    add_inlines(p, block["c"])
                    if (
                        text.startswith("Status:")
                        or text.startswith("Source:")
                        or text.startswith("As of:")
                        or text.startswith("Companion:")
                    ):
                        p.style = doc.styles["BKP Metadata"]
            last_heading_level = None
        elif kind in ("BulletList", "OrderedList"):
            render_list(doc, block, decimal_num_id, bullet_num_id)
            last_heading_level = None
        elif kind == "CodeBlock":
            payload = block["c"][1]
            parsed = None
            if current_component == "DIRECTORY":
                try:
                    candidate = yaml.safe_load(payload)
                except yaml.YAMLError:
                    candidate = None
                if isinstance(candidate, dict) and "status" in candidate:
                    parsed = candidate
            if parsed is not None:
                def open_section(name, _n=[part_number]):
                    _n[0] += 1
                    add_part_opening(doc, name, _n[0],
                                     subtitle=SECTION_SUBTITLE.get(name))
                render_register(doc, parsed, open_section)
                part_number += 2
            else:
                add_code_block(doc, payload)
            last_heading_level = None
        elif kind == "HorizontalRule":
            add_horizontal_rule(doc)
            last_heading_level = None
        elif kind == "Table":
            add_table(doc, block)
            last_heading_level = None

    # CORE and PROCESSES carry their sub-line as body text in the source, so it
    # never reaches add_part_opening's subtitle path and stayed a plain
    # paragraph at the left margin while STATUS and REFS were tucked under
    # their titles. Caught 100826: "processes doesn't". Restyled here, after
    # the body is built, by matching the known sub-lines where they follow a
    # Heading 1 - the parse flow is left alone.
    _subs = {v.lower() for v in SECTION_SUBTITLE.values()}
    _paras = doc.paragraphs
    for _i, _p in enumerate(_paras[:-1]):
        if _p.style.name != "Heading 1":
            continue
        _next = _paras[_i + 1]
        if _next.text.strip().lower() in _subs and _next.style.name != "BKP Part Subtitle":
            _p.paragraph_format.space_after = Pt(0)
            _next.style = doc.styles["BKP Part Subtitle"]
            _next.alignment = _p.alignment
            _next.paragraph_format.left_indent = _p.paragraph_format.left_indent

    # add_contents(doc)  # OFF 100826 - see the docstring. Needs the folio
    # restart ruled on first: a page-numbered index cannot address a volume
    # whose numbering restarts in every part.
    set_update_fields(doc)
    for section_index, section in enumerate(doc.sections):
        pg_num = section._sectPr.find(qn("w:pgNumType"))
        if section_index == 0:
            if pg_num is None:
                pg_num = OxmlElement("w:pgNumType")
                section._sectPr.append(pg_num)
            pg_num.set(qn("w:start"), "1")
        elif pg_num is not None:
            section._sectPr.remove(pg_num)
    doc.core_properties.title = "BKP Pipeline - Compiled Governance Set"
    doc.core_properties.subject = " + ".join(COMPONENTS)
    doc.core_properties.author = "Bangkok Post Desk Editor project"
    doc.core_properties.comments = "Mechanically generated from COMPILED.md"
    # Page numbering restarts in every part except the cover. Done as a
    # post-pass: python-docx clones the sentinel sectPr on add_section(), so
    # mutating at creation time lands on the preceding section.
    for sec in list(doc.sections)[1:]:
        restart_page_numbering(sec, 1)

    # ---- shortlink / version pair -------------------------------------
    # The version line is restyled to match the shortlink above it (Arial,
    # bold, grey) so the two read as one block, and the section rule is hung
    # under the pair rather than under the head.
    slug_re = re.compile(r'^\d{6}_all_records-extracted$')
    for para in doc.paragraphs:
        if slug_re.match(para.text.strip()):
            for r in para.runs:
                r.font.name = "Arial"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                rpr = r._element.get_or_add_rPr()
                for tag in ("w:rFonts",):
                    for el in rpr.findall(qn(tag)):
                        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                            el.set(qn(a), "Arial")
            para.paragraph_format.space_after = Pt(2)
            set_paragraph_bottom_rule(para, 30, 10)

    doc.save(output)
    write_manifest(manifest, source, output, ast, skipped_separators)
    print(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--reference", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--pandoc", type=Path, default=Path(r"C:\Program Files\Pandoc\pandoc.exe"))
    parser.add_argument("--manifest", type=Path)
    parser.add_argument("--component", choices=COMPONENTS, default=None,
                        help="build a single component companion; default is the full compiled build")
    args = parser.parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    manifest = args.manifest or args.output.with_suffix(".manifest.json")
    build(args.source, args.reference, args.output, args.pandoc, manifest,
          component=args.component)


if __name__ == "__main__":
    main()
