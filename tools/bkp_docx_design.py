"""Reusable BKP DOCX design primitives and the original two-page proof builder."""

from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Documents\BANGKOK POST DESK EDITOR\GPT_Mirror")
REFERENCE = ROOT / "compare" / "the_bangkok_post_blueprint_compendium_final.docx"
OUTPUT = ROOT / "compare" / "BKP_conversion_design_proof.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="000000", size=8):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_paragraph_bottom_rule(paragraph, size=18, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def set_paragraph_left_rule(paragraph, size=30, space=10):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), "000000")
    p_bdr.append(left)


def set_repeat_table_layout(table, widths_twips):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    total = sum(widths_twips)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_numbering_definition(doc, fmt, text, left=480, hanging=240):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_page_field(paragraph, bold=False, tab=True, size=None):
    if tab:
        paragraph.add_run("\t")
    run = paragraph.add_run()
    if bold:
        # Hard black folio: it must read as a page number, not as part of the
        # washed-out section slug beside it.
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if size is not None:
        run.font.size = size
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(node)


def clear_document_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_font(style, name, size, bold=False, italic=False, color="000000"):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    normal = doc.styles["normal"]
    set_font(normal, "Arial", 11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_font(h1, "Arial Black", 38, True)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(13)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_font(h2, "Arial Black", 21, True)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(9)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_font(h3, "Arial Black", 13.5, True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    def ensure(name, base="normal"):
        try:
            return doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles[base]
            return style

    metadata = ensure("BKP Metadata")
    set_font(metadata, "Arial", 10.5)
    metadata.paragraph_format.space_after = Pt(5)

    axiom_label = ensure("BKP Axiom Label")
    set_font(axiom_label, "Arial Black", 7.5, True)
    axiom_label.paragraph_format.space_before = Pt(9)
    axiom_label.paragraph_format.space_after = Pt(1)
    axiom_label.paragraph_format.left_indent = Inches(0.16)
    axiom_label.paragraph_format.keep_with_next = True

    axiom = ensure("BKP Axiom")
    set_font(axiom, "Arial", 13, True)
    axiom.paragraph_format.space_after = Pt(10)
    axiom.paragraph_format.left_indent = Inches(0.16)

    note = ensure("BKP Operator Note")
    set_font(note, "Arial", 10.5, False, True)
    note.paragraph_format.left_indent = Inches(0.18)
    note.paragraph_format.space_before = Pt(9)
    note.paragraph_format.space_after = Pt(9)

    register = ensure("BKP Register")
    set_font(register, "Arial Black", 8, True)
    register.paragraph_format.space_after = Pt(0)

    footer = ensure("BKP Footer")
    set_font(footer, "Arial", 8, True, False, "555555")
    footer.paragraph_format.space_after = Pt(0)

    list_style = ensure("BKP Numbered")
    set_font(list_style, "Arial", 10.5)
    list_style.paragraph_format.space_after = Pt(6)
    list_style.paragraph_format.line_spacing = 1.08

    bullet_style = ensure("BKP Bullet")
    set_font(bullet_style, "Arial", 10.5)
    bullet_style.paragraph_format.space_after = Pt(5)
    bullet_style.paragraph_format.line_spacing = 1.08


def add_metadata(doc, label, value):
    p = doc.add_paragraph(style="BKP Metadata")
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(value)
    return p


def add_register(doc):
    table = doc.add_table(rows=1, cols=4)
    set_repeat_table_layout(table, [2340, 2340, 2340, 2340])
    labels = [("01", "BLUEPRINT"), ("02", "PROCESSES"), ("03", "STATUS"), ("04", "REFERENCES")]
    for idx, (num, name) in enumerate(labels):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 55, 90, 55, 90)
        if idx == 0:
            set_cell_shading(cell, "000000")
        p = cell.paragraphs[0]
        p.style = doc.styles["BKP Register"]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(num + "  ")
        r2 = p.add_run(name)
        if idx == 0:
            r1.font.color.rgb = RGBColor(255, 255, 255)
            r2.font.color.rgb = RGBColor(255, 255, 255)
    set_table_borders(table, "000000", 10)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(7)


def add_axiom(doc, text):
    label = doc.add_paragraph("AXIOM", style="BKP Axiom Label")
    set_paragraph_left_rule(label)
    body = doc.add_paragraph(text, style="BKP Axiom")
    set_paragraph_left_rule(body)


def add_header_footer(section, section_name):
    for header in (section.header, section.even_page_header):
        hp = header.paragraphs[0]
        hp.clear()
        hp.style = "BKP Footer"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.add_run(f"BKP GOVERNANCE SET  /  {section_name}")
        set_paragraph_bottom_rule(hp, 6, 2)
    for footer in (section.footer, section.even_page_footer):
        fp = footer.paragraphs[0]
        fp.clear()
        fp.style = "BKP Footer"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.86), WD_TAB_ALIGNMENT.RIGHT)
        fp.add_run("CONTROLLED INSTITUTIONALISM  /  CONVERSION PROOF")
        add_page_field(fp)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    clear_document_body(doc)
    configure_styles(doc)
    doc.settings.odd_and_even_pages_header_footer = True

    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.76)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)
    add_header_footer(section, "BLUEPRINT")
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.")
    bullet_num_id = add_numbering_definition(doc, "bullet", "●")

    add_register(doc)
    h1 = doc.add_paragraph("BLUEPRINT", style="Heading 1")
    set_paragraph_bottom_rule(h1, 30, 6)
    add_metadata(doc, "Status", "Live")
    add_metadata(doc, "Architecture", "Governed discretion / Lean RAG retrieval / Operator-first verification")
    add_metadata(doc, "Companions", "PROCESSES, STATUS, REFERENCES")

    h2 = doc.add_paragraph("OPERATING DOCTRINE & TONE", style="Heading 2")
    set_paragraph_bottom_rule(h2, 14, 4)
    doc.add_paragraph("ORIENTATION", style="Heading 3")
    add_axiom(doc, "Solve the problem.")
    doc.add_paragraph(
        "Every exchange aims at the editorial problem in front of it. Brevity, disagreement and "
        "holding a position under pushback all follow from that - they are not the goal; solving is."
    )
    doc.add_paragraph(
        "The same rule governs the register: act to resolve the copy, not to populate the documents. "
        "A reply or an edit that serves the desk's process rather than the copy has lost the orientation."
    )
    doc.add_paragraph("PRECISION & CHARACTER", style="Heading 3")
    doc.add_paragraph(
        "Marry mechanical precision with editorial character. Craft smart, concise, sometimes witty, "
        "sometimes sombre heads and decks. Heads are statements, qualified by a deck - not explanations."
    )

    doc.add_page_break()

    h2b = doc.add_paragraph("VERIFICATION PROTOCOL", style="Heading 2")
    set_paragraph_bottom_rule(h2b, 14, 4)
    doc.add_paragraph("OPERATOR-FIRST", style="Heading 3")
    add_axiom(doc, "Resolve checks in this order. Higher steps come first.")

    items = [
        ("CHECK STATUS.", "Scan on the first required status check and retain it for the edit. Listed entity -> apply the register; no search."),
        ("ASK THE OPERATOR.", "The operator is in the chair and closes most checks in a line. Do not perform external work to settle a question the operator can resolve instantly."),
        ("LIVE WEB SEARCH.", "Run only on the operator's explicit request or prior consent. A bangkokpost.com site search is the preferred method."),
    ]
    for label, body in items:
        p = doc.add_paragraph(style="BKP Numbered")
        apply_numbering(p, decimal_num_id)
        r = p.add_run(label + " ")
        r.bold = True
        p.add_run(body)

    doc.add_paragraph("RULES OF THE PROTOCOL", style="Heading 3")
    bullets = [
        "Operator ruling overrides all other sources of truth.",
        'A flag never means a change. "Flag" means exactly: ask about it, but don\'t alter.',
        "Filed copy contradicting the register or a confirmed fact is flagged for operator guidance.",
        "Scan before you change.",
    ]
    for text in bullets:
        p = doc.add_paragraph(text, style="BKP Bullet")
        apply_numbering(p, bullet_num_id)

    note = doc.add_paragraph(style="BKP Operator Note")
    set_paragraph_left_rule(note, 18, 8)
    r = note.add_run("OPERATOR NOTE  ")
    r.bold = True
    r.italic = False
    note.add_run(
        "The register governs known status. Silence is not permission to improvise; it returns the decision to the chair."
    )

    set_update_fields(doc)
    doc.core_properties.title = "BKP conversion design proof"
    doc.core_properties.subject = "Two-page mechanical Markdown-to-Word design demonstration"
    doc.core_properties.author = "Bangkok Post Desk Editor project"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
